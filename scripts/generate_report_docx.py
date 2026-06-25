#!/usr/bin/env python3
"""Generate Word document project report from structured content."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "PROJECT_REPORT.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Multi-Agent Research Intelligence System\nProject Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "\nDinaz Sanaullah\n"
        "GCP Project: multi-agent-research-system\n"
        "GitHub: github.com/Dinaz-Sanaullah/multi-agent-research-system\n"
        "June 2025"
    ).font.size = Pt(12)

    doc.add_page_break()

    sections = [
        ("1. Executive Summary", [
            "This project implements a Multi-Agent Research Intelligence System using Google's Agent Development Kit (ADK). "
            "The system combines RAG over academic PDFs with real-time Tavily web search. Seven specialized agents "
            "collaborate through sequential, parallel, hierarchical, and feedback-loop patterns to produce "
            "quality-reviewed, well-cited research responses deployed on Google Cloud Run.",
        ]),
        ("2. Objectives", [
            "Design multi-agent system with Google ADK (Orchestrator, Researcher, Reviewer)",
            "Integrate RAG: PDF extraction, semantic chunking, text-embedding-005, FAISS",
            "Implement Tavily web search with intelligent RAG vs Web routing",
            "Build custom tools for agent composition",
            "Add observability and monitoring",
            "Build Streamlit UI",
            "Deploy to Google Cloud Run",
        ]),
        ("3. Architecture", [
            "Orchestrator (root) → Research Pipeline (SequentialAgent)",
            "Pipeline: Planner → Research Team → Quality Review Loop",
            "Research Team: RAG Researcher + Web Researcher (parallel or sequential)",
            "Quality Loop: Synthesizer → Reviewer → Refiner (LoopAgent)",
            "LLM: Gemini 2.5 Flash | Embeddings: text-embedding-005 | Vector DB: FAISS",
        ]),
        ("4. Multi-Agent System", [
            "Orchestrator — coordinates pipeline, presents final answer",
            "Planner — classifies query, creates research plan",
            "RAG Researcher — searches FAISS knowledge base",
            "Web Researcher — Tavily search for latest citations",
            "Synthesizer — merges findings into coherent response",
            "Reviewer/QA — quality gate with exit_loop approval",
            "Refiner — revises synthesis based on feedback",
        ]),
        ("5. Communication Patterns", [
            "Sequential Flow: Planner → Research → Quality Loop",
            "Parallel Execution: RAG + Web researchers concurrently",
            "Hierarchical Delegation: Orchestrator → sub-pipeline",
            "Feedback Loop: Synthesize → Review → Refine until approved",
        ]),
        ("6. RAG Pipeline", [
            "PDF extraction (PyPDF) → semantic chunking → Vertex AI embeddings → FAISS",
            "Sample corpus: Attention Is All You Need, BERT, GPT-3 (104 chunks)",
            "Retrieval scores: 0.71–0.75 for test queries",
        ]),
        ("7. Web Search", [
            "Tavily API with advanced search and AI answer summary",
            "Intelligent routing: RAG vs Web vs Hybrid based on query signals",
            "Test: 3 results in ~6 seconds for latest transformer papers query",
        ]),
        ("8. Observability", [
            "structlog JSON logging for agent events and tool calls",
            "In-memory metrics: query counts, latency, review iterations",
            "Optional OpenTelemetry Cloud Trace integration",
        ]),
        ("9. Deployment", [
            "Google Cloud Run: research-agent (us-central1)",
            "URL: https://research-agent-chfrwn7sia-uc.a.run.app/",
            "CI/CD via Cloud Build, secrets via Secret Manager",
            "ADK Web UI enabled for production interaction",
        ]),
        ("10. Challenges & Solutions", [
            "429 quota errors → retries + sequential research mode",
            "Model availability → gemini-2.5-flash",
            "Tool registration errors → both tools on both researchers",
            "Embedding limits → batched API calls",
        ]),
        ("11. Conclusion", [
            "All project requirements fulfilled. The system demonstrates production-ready "
            "multi-agent orchestration with RAG, web search, observability, and cloud deployment. "
            "Future work: persistent sessions, larger corpora, and agent evaluation benchmarks.",
        ]),
    ]

    for heading, content in sections:
        add_heading(doc, heading, level=1)
        if len(content) == 1 and len(content[0]) > 100:
            doc.add_paragraph(content[0])
        else:
            add_bullets(doc, content)
        doc.add_paragraph()

    doc.save(str(OUTPUT))
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
